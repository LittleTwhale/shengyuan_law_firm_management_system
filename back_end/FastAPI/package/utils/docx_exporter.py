"""
Markdown → DOCX 转换器
将 DeepSeek 生成的 Markdown 分析报告转换为格式化的 Word 文档
"""
import re
import io
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =================================================================
#  常量配置
# =================================================================

# 字体设置
FONT_BODY = "宋体"
FONT_HEADING = "黑体"
FONT_CODE = "Consolas"

# 字号（Pt）
SIZE_H1 = 22
SIZE_H2 = 18
SIZE_H3 = 16
SIZE_BODY = 12
SIZE_CODE = 10
SIZE_FOOTER = 9

# 颜色
COLOR_H2 = RGBColor(0x16, 0x5D, 0xFF)  # #165DFF 蓝色
COLOR_GRAY = RGBColor(0x99, 0x99, 0x99)
COLOR_DARK = RGBColor(0x33, 0x33, 0x33)


# =================================================================
#  辅助函数
# =================================================================

def _set_cell_shading(cell, color: str):
    """设置表格单元格底色"""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _add_hyperlink(paragraph, text: str, url: str):
    """在段落中添加可点击的超链接"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # 设置蓝色+下划线
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "165DFF")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run_elem.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run_elem.append(t)
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def _add_run(paragraph, text: str, bold=False, italic=False, font_name=None, font_size=None, color=None):
    """在段落中添加一个 run"""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name) if run._element.rPr else None
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color
    return run


# =================================================================
#  Markdown → DOCX 主函数
# =================================================================

def markdown_to_docx(
    markdown: str,
    case_number: str = "",
    case_category: str = "",
) -> io.BytesIO:
    """
    将 Markdown 文本转换为 Word (.docx) 文件的字节流

    Args:
        markdown: Markdown 格式的报告内容
        case_number: 案号（用于页眉）
        case_category: 案件类别

    Returns:
        BytesIO 对象（可直接用于 StreamingResponse）
    """
    doc = Document()

    # ========== 页面设置 ==========
    section = doc.sections[0]
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ========== 页眉 ==========
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_text = f"案件智能分析报告"
    if case_number:
        header_text += f" | {case_number}"
    _add_run(header_para, header_text, font_name=FONT_BODY, font_size=SIZE_FOOTER, color=COLOR_GRAY)

    # ========== 页脚（页码） ==========
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(footer_para, "第 ", font_name=FONT_BODY, font_size=SIZE_FOOTER, color=COLOR_GRAY)
    # 插入页码域
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run_pg = footer_para.add_run()
    run_pg._element.append(fldChar1)
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run_instr = footer_para.add_run()
    run_instr._element.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run_end = footer_para.add_run()
    run_end._element.append(fldChar2)
    _add_run(footer_para, f" / 共 ", font_name=FONT_BODY, font_size=SIZE_FOOTER, color=COLOR_GRAY)
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    run_np = footer_para.add_run()
    run_np._element.append(fldChar3)
    instrText2 = OxmlElement("w:instrText")
    instrText2.set(qn("xml:space"), "preserve")
    instrText2.text = " NUMPAGES "
    run_instr2 = footer_para.add_run()
    run_instr2._element.append(instrText2)
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run_end2 = footer_para.add_run()
    run_end2._element.append(fldChar4)
    _add_run(footer_para, " 页", font_name=FONT_BODY, font_size=SIZE_FOOTER, color=COLOR_GRAY)

    # ========== 报告标题 ==========
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(title_para, "案件智能分析报告", bold=True, font_name=FONT_HEADING, font_size=SIZE_H1, color=COLOR_DARK)

    # ========== 元信息行 ==========
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_parts = []
    if case_number:
        meta_parts.append(f"案号：{case_number}")
    if case_category:
        meta_parts.append(f"案件类别：{case_category}")
    _add_run(meta_para, " | ".join(meta_parts), font_name=FONT_BODY, font_size=SIZE_FOOTER, color=COLOR_GRAY)

    # 分隔线
    doc.add_paragraph()  # 空行

    # ========== 渲染 Markdown 正文 ==========
    lines = markdown.split("\n")
    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # -- 代码块 --
        if line.strip().startswith("```"):
            if in_code_block:
                # 结束代码块
                code_text = "\n".join(code_buffer)
                p = doc.add_paragraph()
                p.paragraph_format.keep_with_next = True
                for cl in code_text.split("\n"):
                    _add_run(p, cl + "\n", font_name=FONT_CODE, font_size=SIZE_CODE)
                # 灰色底纹
                from docx.oxml import parse_xml
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F5F5F5")
                shd.set(qn("w:val"), "clear")
                pPr.append(shd)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # -- 空行 --
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # -- 标题 --
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            for run in p.runs:
                run.font.name = FONT_HEADING
                run.font.size = Pt(SIZE_H1)
                run.font.color.rgb = COLOR_DARK
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            for run in p.runs:
                run.font.name = FONT_HEADING
                run.font.size = Pt(SIZE_H2)
                run.font.color.rgb = COLOR_H2
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            for run in p.runs:
                run.font.name = FONT_HEADING
                run.font.size = Pt(SIZE_H3)
                run.font.color.rgb = COLOR_DARK
            i += 1
            continue

        if line.startswith("#### "):
            p = doc.add_heading(line[5:].strip(), level=4)
            for run in p.runs:
                run.font.name = FONT_HEADING
                run.font.size = Pt(SIZE_BODY)
                run.bold = True
            i += 1
            continue

        # -- 有序列表 --
        ol_match = re.match(r"^\d+[.)]\s+(.*)", line)
        if ol_match:
            p = doc.add_paragraph(style="List Number")
            p.clear()
            _render_inline(p, ol_match.group(1))
            i += 1
            continue

        # -- 无序列表 --
        ul_match = re.match(r"^[-*]\s+(.*)", line)
        if ul_match:
            p = doc.add_paragraph(style="List Bullet")
            p.clear()
            _render_inline(p, ul_match.group(1))
            i += 1
            continue

        # -- GFM Task List --
        task_match = re.match(r"^- \[([ xX])\]\s+(.*)", line)
        if task_match:
            checked = task_match.group(1).lower() == "x"
            text = task_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            prefix = "☑ " if checked else "☐ "
            _add_run(p, prefix, font_name=FONT_BODY, font_size=SIZE_BODY)
            _render_inline(p, text)
            i += 1
            continue

        # -- 水平线 --
        if re.match(r"^---+\s*$", line):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # -- 引用块 --
        if line.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].startswith(">"):
                quote_lines.append(lines[i][1:].strip())
                i += 1
            quote_text = "\n".join(quote_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            # 左侧粗边框
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "12")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), "E6A23C")
            pBdr.append(left)
            pPr.append(pBdr)
            # 浅灰底色
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "FDF6EC")
            shd.set(qn("w:val"), "clear")
            pPr.append(shd)
            _render_inline(p, quote_text)
            continue

        # -- 表格 --
        if "|" in line and line.count("|") >= 3:
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].count("|") >= 2:
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:  # 至少标题行 + 分隔行
                _render_table(doc, table_lines)
            continue

        # -- 普通段落 --
        p = doc.add_paragraph()
        _render_inline(p, line)
        p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
        i += 1

    # ========== 免责声明 ==========
    doc.add_paragraph()
    disclaimer_p = doc.add_paragraph()
    disclaimer_p.paragraph_format.left_indent = Cm(1.0)
    # 灰色底框
    pPr = disclaimer_p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "FAFAFA")
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), "CCCCCC")
    pBdr.append(top)
    pPr.append(pBdr)
    _add_run(disclaimer_p, "免责声明：", bold=True, font_name=FONT_HEADING, font_size=SIZE_BODY, color=RGBColor(0x85, 0x64, 0x04))
    _add_run(disclaimer_p, "本报告由 AI 自动生成，仅供律师参考，不构成正式法律意见。报告中涉及的法条引用、策略建议等，请结合专业判断和实际情况审慎使用。",
             font_name=FONT_BODY, font_size=SIZE_BODY, color=RGBColor(0x85, 0x64, 0x04))

    # ========== 写入字节流 ==========
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _render_inline(paragraph, text: str):
    """渲染行内 Markdown（加粗、行内代码、链接）"""
    # 处理行内代码 `code`
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            _add_run(paragraph, part[1:-1], font_name=FONT_CODE, font_size=SIZE_CODE)
        else:
            # 处理 **加粗**
            bold_parts = re.split(r"(\*\*[^*]+\*\*)", part)
            for bp in bold_parts:
                if bp.startswith("**") and bp.endswith("**"):
                    _add_run(paragraph, bp[2:-2], bold=True, font_name=FONT_BODY, font_size=SIZE_BODY)
                else:
                    # 处理 [链接文本](url)
                    link_parts = re.split(r"(\[([^\]]+)\]\(([^)]+)\))", bp)
                    j = 0
                    while j < len(link_parts):
                        if j + 3 < len(link_parts) and link_parts[j+1].startswith("[") and link_parts[j+1].endswith(")" ):
                            # 匹配到完整链接
                            text_part = re.match(r"\[([^\]]+)\]\(([^)]+)\)", link_parts[j+1])
                            if text_part:
                                _add_hyperlink(paragraph, text_part.group(1), text_part.group(2))
                            j += 4
                        else:
                            _add_run(paragraph, link_parts[j], font_name=FONT_BODY, font_size=SIZE_BODY)
                            j += 1


def _render_table(doc, table_lines: list[str]):
    """渲染 Markdown 表格"""
    # 解析标题行
    header_cells = [c.strip() for c in table_lines[0].split("|") if c.strip()]
    # 解析数据行（从第2行开始，跳过分隔行）
    data_rows = []
    for line in table_lines[2:]:
        cells = [c.strip() for c in line.split("|") if c.strip()]
        if cells:
            data_rows.append(cells)

    if not header_cells:
        return

    # 创建表格
    table = doc.add_table(rows=1 + len(data_rows), cols=len(header_cells))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表头行
    for j, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        cell.text = ""
        _add_run(cell.paragraphs[0], cell_text, bold=True, font_name=FONT_HEADING, font_size=SIZE_BODY)
        _set_cell_shading(cell, "F2F2F2")

    # 数据行
    for i, row_cells in enumerate(data_rows):
        for j, cell_text in enumerate(row_cells):
            if j < len(header_cells):
                cell = table.rows[i + 1].cells[j]
                cell.text = ""
                _add_run(cell.paragraphs[0], cell_text, font_name=FONT_BODY, font_size=SIZE_BODY)
