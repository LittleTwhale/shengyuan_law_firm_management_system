# utils/ocr_helper.py
import logging
import os
import cv2
import numpy as np

# PDF与Word处理
import pdfplumber
from docx import Document

from rapidocr_onnxruntime import RapidOCR

logger = logging.getLogger(__name__)

# ============================================================
# 初始化 RapidOCR 引擎 (基于 ONNX Runtime, 稳定无 Bug)
# ============================================================
ocr_engine = None

print("--- 正在初始化 RapidOCR ---")
try:
    ocr_engine = RapidOCR()
    print("--- RapidOCR 初始化成功 ---")
    logger.info("RapidOCR initialized successfully.")
except Exception as e:
    print(f"!!! RapidOCR 初始化出错: {e}")
    ocr_engine = None


def extract_text_from_docx(file_path: str) -> str:
    """提取 Word (.docx) 文本（包含段落、表格，以及内嵌图片的OCR识别）"""
    try:
        doc = Document(file_path)
        full_text = []

        # 1. 提取普通段落文本
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)

        # 2. 提取表格内的文本
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    if cell_text and cell_text not in row_texts:
                        row_texts.append(cell_text)
                if row_texts:
                    full_text.append(" | ".join(row_texts))

        # 3. 提取文档中内嵌的图片，并进行 OCR 识别
        # 遍历文档的所有关联部件 (relationships)
        image_index = 1
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    # 获取图片的二进制数据
                    image_bytes = rel.target_part.blob
                    # 将二进制数据转为 numpy 数组 (OpenCV 格式)
                    nparr = np.frombuffer(image_bytes, np.uint8)
                    img_np = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

                    if img_np is not None:
                        # 复用你现有的图片 OCR 函数 (传入 is_bgr=True)
                        ocr_text = _ocr_image_data(img_np, is_bgr=True)
                        if ocr_text:
                            full_text.append(f"--- [内嵌图片 {image_index} 识别内容] ---")
                            full_text.append(ocr_text)
                            image_index += 1
                except Exception as img_e:
                    logger.warning(f"Word 内嵌图片识别跳过: {img_e}")
                    continue

        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Docx error: {e}")
        return ""


def _read_image_robust(file_path: str):
    """读取图片：解决 Windows 中文路径问题"""
    try:
        if not os.path.exists(file_path):
            return None
        return cv2.imdecode(np.fromfile(file_path, dtype=np.uint8), cv2.IMREAD_COLOR)
    except Exception as e:
        logger.error(f"Image read error: {e}")
        return None


def _ocr_image_data(image_data, is_bgr: bool = False) -> str:
    """
    对图片数据进行 RapidOCR 识别。
    :param image_data: numpy 数组
    :param is_bgr: True 表示 image_data 是 BGR 格式（OpenCV 来源）
    """
    if ocr_engine is None:
        return "[错误] OCR模型未成功启动，请联系管理员"

    try:
        # RapidOCR 官方推荐输入 BGR 格式 (OpenCV 默认格式)
        # 如果来源是 PIL/pdfplumber (RGB格式)，我们需要把它转为 BGR
        if not is_bgr and image_data is not None and len(image_data.shape) == 3:
            if image_data.shape[2] == 3:
                image_data = cv2.cvtColor(image_data, cv2.COLOR_RGB2BGR)

        # 核心识别调用，返回结果和耗时
        result, elapse = ocr_engine(image_data)

        if not result:
            return ""

        # RapidOCR 返回结构: [ [ [[x,y],...], '识别出的文本', 0.99 (置信度) ], ... ]
        page_text = []
        for line in result:
            if len(line) >= 2:
                text = line[1]  # 索引1就是文本内容
                if text and str(text).strip():
                    page_text.append(str(text).strip())

        return "\n".join(page_text)

    except Exception as e:
        error_msg = f"[OCR运行时错误: {str(e)}]"
        logger.error(error_msg)
        print(f"!!! {error_msg}")
        return ""


def extract_pdf_hybrid(file_path: str, min_text_len: int = 50) -> str:
    """
    PDF 混合提取策略
    """
    full_content = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                # 1. 尝试直接提取文本
                text = page.extract_text() or ""
                clean_text = text.replace(" ", "").replace("\n", "")

                if len(clean_text) > min_text_len:
                    full_content.append(f"--- 第 {i + 1} 页 (电子提取) ---\n{text}")
                else:
                    #  扫描件 -> 图片 -> OCR
                    # 200dpi 够用且快
                    pil_image = page.to_image(resolution=200).original
                    # PIL Image → numpy array (RGB)
                    img_array = np.array(pil_image)
                    # 传入 is_bgr=False，因为 PIL 生成的是 RGB，上面的函数会自动处理
                    ocr_text = _ocr_image_data(img_array, is_bgr=False)
                    full_content.append(f"--- 第 {i + 1} 页 (OCR识别) ---\n{ocr_text}")

    except Exception as e:
        logger.error(f"PDF Hybrid Extract Error: {e}")
        return f"PDF处理错误: {str(e)}"

    return "\n\n".join(full_content)


def perform_smart_extraction(file_path: str, file_type: str) -> str:
    """
    统一入口：智能识别文件内容
    """
    if not os.path.exists(file_path):
        return "文件路径不存在"

    file_ext = os.path.splitext(file_path)[1].lower()

    try:
        # 1. Word (.docx)
        if file_ext == '.docx':
            return extract_text_from_docx(file_path)

        # 2. Word (.doc)
        if file_ext == '.doc':
            pdf_path = os.path.splitext(file_path)[0] + ".pdf"
            if os.path.exists(pdf_path):
                return extract_pdf_hybrid(pdf_path)
            else:
                return "待转换PDF文件未找到"

        # 3. PDF
        if file_ext == '.pdf' or file_type == 'application/pdf':
            return extract_pdf_hybrid(file_path)

        # 4. 图片
        if file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tif', '.webp']:
            img_array = _read_image_robust(file_path)
            if img_array is not None:
                return _ocr_image_data(img_array, is_bgr=True)
            else:
                return "图片读取失败"
    except Exception as e:
        logger.error(f"Extraction execution failed: {e}")
        return ""

    return ""