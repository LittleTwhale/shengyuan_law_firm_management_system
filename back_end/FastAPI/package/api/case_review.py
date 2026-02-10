# api/case_review.py
import os
import tempfile
from datetime import datetime

from fastapi.responses import FileResponse
from fastapi import BackgroundTasks
from docx import Document
from ..models.case import Case
from fastapi import APIRouter, Depends, Query, HTTPException, status
from sqlalchemy.orm import Session
from typing import Optional

from ..database.database import get_db
from ..schemas.case import CasePageOut, CaseSimpleOut, CaseOut
from ..crud.case_review import list_pending_cases, count_pending_cases, update_review_status, \
    check_interest_conflict_for_case, replace_text_in_paragraph

from ..core.config import TEMPLATE_DIR

router = APIRouter(
    prefix="/case_review",
    tags=["case_review"]
)


@router.get("/pending", response_model=CasePageOut)
def get_pending_cases(
        role: Optional[str] = None,
        skip: int = Query(0, ge=0),
        limit: int = Query(10, ge=1, le=100),
        db: Session = Depends(get_db)
):
    """
    获取待审核案件列表（仅管理员可访问）
    """
    # 验证管理员权限
    if not role or role not in ["admin", "owner"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="无审核案件权限"
        )

    cases = list_pending_cases(db, skip=skip, limit=limit)
    total = count_pending_cases(db)
    cases_simple = [CaseSimpleOut.model_validate(case) for case in cases]
    return {"items": cases_simple, "total": total}


@router.put("/{case_id}/review", response_model=CaseOut)
def review_case(
        case_id: int,
        reviewer_id: int,
        review_status: str,
        role: Optional[str] = None,
        force: bool = Query(False, description="是否强制通过（忽略利益冲突）"),  # 新增参数
        db: Session = Depends(get_db)
):
    """
    审核案件（通过/拒绝，仅管理员可操作）
    """
    if not role or role not in ["admin", "owner"]:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="无审核案件权限")

    # 1. 审核通过逻辑
    if review_status == "已审核":
        # 如果不是强制通过，则进行冲突检测
        if not force:
            conflict_result = check_interest_conflict_for_case(db, case_id)
            if conflict_result["has_conflict"]:
                # 409 Conflict: 返回详细信息给前端展示
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail={
                        "code": "INTEREST_CONFLICT",
                        "message": "检测到潜在利益冲突，是否强制通过？",
                        "conflicts": conflict_result["details"]
                    }
                )

    # 2. 执行更新
    try:
        updated_case = update_review_status(
            db=db,
            case_id=case_id,
            review_status=review_status,
            reviewer_id=reviewer_id
        )
        if not updated_case:
            raise HTTPException(status_code=404, detail="案件不存在")

        return updated_case

    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

@router.put("/{case_id}/force_review", response_model=CaseOut)
def review_case(
        case_id: int,
        reviewer_id: int,
        review_status: str,  # 接收"已审核"或"已拒绝"
        role: Optional[str] = None,
        db: Session = Depends(get_db)
):
    """
    案件强制通过（不检测案件冲突）
    """
    # 验证管理员权限
    if not role or role not in ["admin", "owner"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="无审核案件权限"
        )

    try:
        updated_case = update_review_status(
            db=db,
            case_id=case_id,
            review_status=review_status,
            reviewer_id=reviewer_id
        )
        if not updated_case:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="案件不存在或已被删除"
            )
        return updated_case
    except ValueError as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        )


@router.get("/{case_id}/approval_form", response_class=FileResponse)
def generate_approval_form(
    case_id: int,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db)
):
    """
    生成并下载案件审批表 (Word格式)
    """
    # 1. 查询案件信息
    case = db.query(Case).filter(Case.case_id == case_id).first()
    if not case:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="案件不存在"
        )

    # 2. 校验状态 (仅已审核通过的案件可以生成审批表)
    # 如果您希望"待审核"状态也能预览，可以注释掉这段
    if case.review_status != "已审核":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"当前案件状态为'{case.review_status}'，仅'已审核'案件可生成审批表"
        )

    # 3. 准备模板路径
    template_path = os.path.join(TEMPLATE_DIR, "case_approval_template.docx")
    if not os.path.exists(template_path):
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="服务器缺少审批表模板文件"
        )

    # 4. 准备填充数据
    # 获取律师真实姓名，如果关联对象为空则给空字符串
    main_lawyer_name = case.main_lawyer.real_name if case.main_lawyer else ""
    assistant_lawyer_name = case.assistant_lawyer.real_name if case.assistant_lawyer else ""


    # 获取导出时间
    export_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    context = {
        "{{case_number}}": case.case_number or "",
        "{{commission_date}}": str(case.commission_date) if case.commission_date else "",
        "{{client_name}}": case.client_name or "",
        "{{client_phone}}": case.client_phone or "",
        "{{client_id_number}}": case.client_id_number or "",
        "{{plaintiff}}": case.plaintiff or "",
        "{{defendant}}": case.defendant or "",
        "{{appellant_info}}":case.appellant_info or "",
        "{{extra_appellant_info}}":case.extra_appellant_info or "",
        "{{court}}": case.court or "",
        "{{case_category}}": case.case_category or "",
        "{{cause}}": case.cause or "",
        "{{main_lawyer_name}}": main_lawyer_name,
        "{{assistant_lawyer_name}}": assistant_lawyer_name,
        "{{fee_method}}": case.fee_method or "",
        "{{case_income}}": str(case.case_income) if case.case_income is not None else "0.00",
        "{{details}}": case.details or "无",
        "{{review_status}}": case.review_status or "",
        "{{reviewer_name}}": case.reviewer.real_name or "",
        "{{export_time}}": export_time,
    }

    try:
        doc = Document(template_path)

        # 5. 执行替换逻辑 (包含段落和表格)
        # 替换段落中的文本
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, context)

        # 替换表格中的文本 (遍历所有表格、所有行、所有单元格)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, context)

        # 6. 保存到临时文件
        # 使用 tempfile 创建临时文件，避免污染服务器目录
        # delete=False 确保文件在关闭后暂时保留供下载，下载后由 background_tasks 删除
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        # 7. 设置下载文件名 (使用 URL 编码避免中文乱码问题，FastAPI会自动处理)
        filename = f"案件审批表_{case.case_number}.docx"

        # 添加后台任务：响应发送后删除临时文件
        background_tasks.add_task(os.remove, tmp_path)

        return FileResponse(
            path=tmp_path,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        # 如果生成过程中出错，记录日志并返回错误
        print(f"Generate document error: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="生成审批表失败，请联系管理员"
        )